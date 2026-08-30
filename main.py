"""
Analisis Rapor Pendidikan v12.0 — API + Sidebar HTML + SQLite DB.
"""
import io
import json
import os
import re
import sqlite3
from datetime import datetime
from typing import List, Optional

import pandas as pd
import pdfplumber
from docx import Document
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

ALLOWED_EXT = {".docx", ".pdf", ".xlsx", ".xls"}
COLUMN_MAP = {
    "no": ["no", "nomor", "#"],
    "capaian": ["capaian", "indikator", "nama indikator", "aspek", "kompetensi"],
    "skor_2025": ["skor rapor 2025", "skor 2025", "nilai 2025"],
    "skor_2024": ["skor rapor 2024", "skor 2024", "nilai 2024"],
    "definisi": ["definisi capaian", "definisi", "pengertian", "deskripsi"],
    "kategori": ["kategori", "kualifikasi", "predikat", "status capaian"],
    "perubahan": ["perubahan skor", "perubahan", "selisih", "delta"],
}

app = FastAPI(title="Analisis Rapor Pendidikan", version="12.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
# ============ MAPPING KODE BENAHI & ARKAS ============
KODE_BENAHI_MAP = {
    "A.1": {
        "akar": "Kompetensi membaca dan analisis teks peserta didik masih rendah",
        "benahi": [
            {
                "kode_pbd": "PBD-LIT-01",
                "kegiatan_pbd": "Peningkatan kompetensi guru dalam literasi membaca melalui Komunitas Belajar (Kombel)",
                "kode_arkas": "03.02.01",
                "kegiatan_arkas": "Pengembangan KKG/MGMP atau Gugus Sekolah"
            },
            {
                "kode_pbd": "PBD-LIT-02",
                "kegiatan_pbd": "Penyediaan bahan bacaan pengayaan untuk mendukung gerakan literasi sekolah",
                "kode_arkas": "05.03.02",
                "kegiatan_arkas": "Pengadaan Buku Teks Utama/Pendamping/Bacaan"
            }
        ]
    },
    "A.2": {
        "akar": "Pemahaman konsep dasar matematika dan pemecahan masalah belum optimal",
        "benahi": [
            {
                "kode_pbd": "PBD-NUM-01",
                "kegiatan_pbd": "Pelatihan pembelajaran numerasi berbasis media konkret dan kontekstual",
                "kode_arkas": "03.02.04",
                "kegiatan_arkas": "Peningkatan Kualitas Guru Mata Pelajaran/Kelas"
            }
        ]
    },
    "D.4": {
        "akar": "Praktik pembelajaran interaktif dan diferensiasi belum berjalan konsisten",
        "benahi": [
            {
                "kode_pbd": "PBD-PBM-01",
                "kegiatan_pbd": "Supervisi akademik dan diskusi peer-teaching antar guru",
                "kode_arkas": "03.01.03",
                "kegiatan_arkas": "Pelaksanaan Supervisi / Evaluasi Pembelajaran"
            }
        ]
    },
    "D.8": {
        "akar": "Penerapan iklim keamanan, pencegahan perundungan, dan inklusivitas belum maksimal",
        "benahi": [
            {
                "kode_pbd": "PBD-IKL-01",
                "kegiatan_pbd": "Sosialisasi dan pembentukan Tim Pencegahan dan Penanganan Kekerasan (TPPK)",
                "kode_arkas": "06.07.01",
                "kegiatan_arkas": "Penyelenggaraan Sekolah Sehat, Aman, Inklusif"
            }
        ]
    }
}
# ============ DATABASE SETUP ============
DB_NAME = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rapor.db")

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    
    # 1. Buat tabel jika belum ada
    query_create = (
        "CREATE TABLE IF NOT EXISTS history ("
        "id INTEGER PRIMARY KEY AUTOINCREMENT, "
        "nama_sekolah TEXT NOT NULL, "
        "tahun TEXT DEFAULT '-', "
        "filename TEXT DEFAULT '', "
        "total_indikator INTEGER DEFAULT 0, "
        "rata_rata_skor REAL DEFAULT 0.0, "
        "indikator_naik INTEGER DEFAULT 0, "
        "indikator_turun INTEGER DEFAULT 0, "
        "jumlah_baik INTEGER DEFAULT 0, "
        "jumlah_sedang INTEGER DEFAULT 0, "
        "jumlah_kurang INTEGER DEFAULT 0, "
        "data JSON NOT NULL, "
        "created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP"
        ")"
    )
    c.execute(query_create)
    
    # 2. Tambahkan kolom secara otomatis jika menggunakan database lama
    c.execute("PRAGMA table_info(history)")
    existing_columns = [col[1] for col in c.fetchall()]
    
    columns_to_add = [
        ("total_indikator", "INTEGER DEFAULT 0"),
        ("rata_rata_skor", "REAL DEFAULT 0.0"),
        ("indikator_naik", "INTEGER DEFAULT 0"),
        ("indikator_turun", "INTEGER DEFAULT 0"),
        ("jumlah_baik", "INTEGER DEFAULT 0"),
        ("jumlah_sedang", "INTEGER DEFAULT 0"),
        ("jumlah_kurang", "INTEGER DEFAULT 0"),
    ]
    
    for col_name, col_type in columns_to_add:
        if col_name not in existing_columns:
            c.execute(f"ALTER TABLE history ADD COLUMN {col_name} {col_type}")
            
    conn.commit()
    conn.close()
    print(f"OK: Database initialized and updated at {DB_NAME}")

# Jalankan inisialisasi database
init_db()

# ============ MODELS ============
class Indikator(BaseModel):
    no: str
    capaian: str
    skor_2025: Optional[float] = None
    definisi_capaian: str = ""
    perubahan_skor: Optional[float] = None
    sub_indikator: List["Indikator"] = []

class RaporResponse(BaseModel):
    judul: str = ""
    capaian: List[Indikator] = []
    baik: List[Indikator] = []
    sedang: List[Indikator] = []
    kurang: List[Indikator] = []
    ringkasan: dict = {}
    kesimpulan: str = ""

# ============ HELPERS ============
def _clean_angka(v):
    if v is None:
        return None
    s = str(v).strip().replace("%", "").replace("Rp", "")
    s = re.sub(r"[^\d.,\-]", "", s)
    if not s or s in (".", "-", ","):
        return None
    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None

def _parse_perubahan(v):
    if v is None:
        return None
    s = str(v).strip().lower()
    if "tidak berubah" in s or "tetap" in s:
        return 0.0
    angka = _clean_angka(s)
    if angka is None:
        return None
    if "turun" in s or "menurun" in s:
        return -angka
    return angka

def _ekstrak_meta(judul, nama_file):
    teks = judul if judul else nama_file
    m = re.search(r"(20\d{2})", teks)
    tahun = m.group(1) if m else ""
    bersih = re.sub(r"(20\d{2})", " ", teks)
    generik = [
        "laporan", "rapor", "pendidikan", "tahun", "pbd", "unduhan",
        "satuan", "perencanaan", "berbasis", "data", "lembar", "dokumen",
    ]
    for g in generik:
        bersih = re.sub(r"\b" + g + r"\b", " ", bersih, flags=re.IGNORECASE)
    bersih = re.sub(r"\s+", " ", bersih).strip("-").strip()
    return {"nama_sekolah": bersih if bersih else "-", "tahun": tahun if tahun else "-"}

def _find_header_row(df):
    for i, row in df.iterrows():
        teks = " ".join(str(c).lower() for c in row.tolist())
        if ("skor" in teks or "capaian" in teks or "indikator" in teks) and (
            "2025" in teks or "2024" in teks or "definisi" in teks
        ):
            return i
    return None

def _normalize(df):
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    rename = {}
    for baku, sinonim in COLUMN_MAP.items():
        for s in sinonim:
            for col in df.columns:
                cl = col.lower()
                if cl == s or (len(s) >= 5 and s in cl):
                    rename[col] = baku
                    break
    return df.rename(columns=rename)

# ============ PARSERS ============
def parse_excel(content):
    xls = pd.ExcelFile(io.BytesIO(content))
    target = None
    for s in xls.sheet_names:
        if "laporan rapor" in s.lower():
            target = s
            break
    if target is None:
        best, best_count = None, -1
        for s in xls.sheet_names:
            n = xls.parse(s, header=None).shape[0]
            if n > best_count:
                best, best_count = s, n
        target = best
    raw = xls.parse(target, header=None)
    judul = ""
    for i in range(min(5, len(raw))):
        for val in raw.iloc[i].tolist():
            if val is not None and "rapor" in str(val).lower() and "pendidikan" in str(val).lower():
                if re.search(r"(20\d{2})", str(val)):
                    judul = str(val).strip()
                    break
        if judul:
            break
    hdr_idx = None
    for i, row in raw.iterrows():
        c0 = str(row.iloc[0]).strip().lower() if len(row) > 0 else ""
        c1 = str(row.iloc[1]).strip().lower() if len(row) > 1 else ""
        if c0 == "no" and ("indikator" in c1 or "capaian" in c1):
            hdr_idx = i
            break
    if hdr_idx is not None:
        df = raw.iloc[hdr_idx + 1:].copy()
        pos_map = {0: "no", 1: "capaian", 2: "kategori", 3: "skor_2025", 4: "definisi", 5: "perubahan", 6: "skor_2024"}
        rename = {df.columns[pos]: baku for pos, baku in pos_map.items() if pos < len(df.columns)}
        df = df.rename(columns=rename)
    else:
        hdr = _find_header_row(raw)
        if hdr is not None:
            df = raw.iloc[hdr + 1:].copy()
            df.columns = [str(c).strip() for c in raw.iloc[hdr].tolist()]
        else:
            df = raw
        df = _normalize(df)
    df = df.dropna(how="all").dropna(axis=1, how="all")
    return df, judul

def parse_pdf(content):
    semua = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for tbl in page.extract_tables():
                for row in tbl:
                    semua.append([c if c is not None else "" for c in row])
    if not semua:
        raise HTTPException(422, "Tidak ditemukan tabel di PDF.")
    df = pd.DataFrame(semua)
    hdr = _find_header_row(df)
    if hdr is not None:
        df.columns = [str(c).strip() for c in df.iloc[hdr].tolist()]
        df = df.iloc[hdr + 1:].reset_index(drop=True)
    return _normalize(df.dropna(how="all"))

def parse_docx(content):
    doc = Document(io.BytesIO(content))
    if not doc.tables:
        raise HTTPException(422, "Tidak ditemukan tabel di dokumen Word.")
    semua = [[cell.text.strip() for cell in row.cells] for tbl in doc.tables for row in tbl.rows]
    df = pd.DataFrame(semua)
    hdr = _find_header_row(df)
    if hdr is not None:
        df.columns = [str(c).strip() for c in df.iloc[hdr].tolist()]
        df = df.iloc[hdr + 1:].reset_index(drop=True)
    return _normalize(df.dropna(how="all"))

# ============ ANALISIS ============
def _analisis(df, judul=""):
    if "capaian" not in df.columns:
        raise HTTPException(422, f"Kolom 'capaian' tidak ditemukan. Kolom: {list(df.columns)}")
    df = df.dropna(subset=["capaian"])
    if df.empty:
        raise HTTPException(422, "Tidak ada baris data yang bisa dianalisis.")

    rows = []
    for _, r in df.iterrows():
        skor25 = _clean_angka(r.get("skor_2025"))
        skor24 = _clean_angka(r.get("skor_2024"))
        perubahan = _parse_perubahan(r.get("perubahan"))
        if perubahan is None and skor25 is not None and skor24 is not None:
            perubahan = round(skor25 - skor24, 2)
        kode = str(r.get("no", "")).strip()
        if not kode or kode.lower() in ("nan", "none"):
            kode = "-"
        nama = str(r.get("capaian", "")).strip()
        definisi = str(r.get("definisi", "")).strip()
        if "\n" in nama:
            bagian = [b.strip() for b in nama.split("\n") if b.strip()]
            nama = bagian[0]
            if not definisi and len(bagian) > 1:
                definisi = " ".join(bagian[1:])
        kateg = str(r.get("kategori", "")).strip()
        if not kateg or kateg.lower() in ("nan", "none", "-"):
            kateg = ""
        rows.append({"no": kode, "capaian": nama, "skor_2025": skor25, "definisi_capaian": definisi, "perubahan_skor": perubahan, "kategori": kateg})

    def norm_kategori(k):
        kl = k.lower()
        if "capaian" in kl: return "Capaian"
        if "baik" in kl or "tinggi" in kl: return "Baik"
        if "sedang" in kl or "cukup" in kl: return "Sedang"
        if "kurang" in kl or "rendah" in kl or "intervensi" in kl: return "Kurang"
        return k.title()

    grouped = {"Capaian": [], "Baik": [], "Sedang": [], "Kurang": []}
    parent_aktif = None
    skor_utama = []
    naik, turun, jumlah_utama = 0, 0, 0
    terbaik, terendah = None, None
    kurang_list = []

    for r in rows:
        kateg = r["kategori"].lower()
        is_utama = any(k in kateg for k in ("capaian", "baik", "sedang", "kurang"))
        item = {"no": r["no"], "capaian": r["capaian"], "skor_2025": r["skor_2025"], "definisi_capaian": r["definisi_capaian"], "perubahan_skor": r["perubahan_skor"], "sub_indikator": []}
        
        if is_utama:
            g = norm_kategori(kateg)
            if g not in grouped: g = "Capaian"
            grouped[g].append(item)
            if g == "Kurang": kurang_list.append((item["capaian"], item["skor_2025"]))
            parent_aktif = item
            jumlah_utama += 1
            if item["skor_2025"] is not None:
                skor_utama.append(item["skor_2025"])
                if terbaik is None or item["skor_2025"] > terbaik[1]: terbaik = (item["capaian"], item["skor_2025"])
                if terendah is None or item["skor_2025"] < terendah[1]: terendah = (item["capaian"], item["skor_2025"])
            if (item["perubahan_skor"] or 0) > 0: naik += 1
            elif (item["perubahan_skor"] or 0) < 0: turun += 1
        else:
            if parent_aktif is not None: parent_aktif["sub_indikator"].append(item)
            else:
                grouped["Capaian"].append(item)
                jumlah_utama += 1

    jpk = {k: len(v) for k, v in grouped.items()}
    rata2 = round(sum(skor_utama) / len(skor_utama), 2) if skor_utama else None

    kesimpulan = (
        f"Berdasarkan hasil analisis Rapor Pendidikan, dari {jumlah_utama} indikator utama "
        f"yang dianalisis, sebanyak {jpk['Baik']} indikator berada pada predikat Baik, "
        f"{jpk['Sedang']} indikator pada predikat Sedang, {jpk['Kurang']} indikator pada "
        f"predikat Kurang, dan {jpk['Capaian']} indikator pada predikat Capaian. "
        f"Rata-rata skor capaian tahun 2025 adalah {rata2 if rata2 is not None else '-'}, "
        f"dengan {naik} indikator mengalami peningkatan skor dan {turun} indikator mengalami "
        f"penurunan dibandingkan tahun 2024."
    )
    if terbaik is not None or terendah is not None:
        kesimpulan += " Capaian tertinggi"
        if terbaik is not None: kesimpulan += f" terdapat pada indikator {terbaik[0]} dengan skor {terbaik[1]}"
        if terbaik is not None and terendah is not None: kesimpulan += ","
        if terendah is not None: kesimpulan += f" sedangkan capaian terendah pada indikator {terendah[0]} dengan skor {terendah[1]}"
        kesimpulan += "."
    if jpk["Kurang"] > 0:
        kesimpulan += (
            f" Terdapat {jpk['Kurang']} indikator yang masih berpredikat "
            f'<span class="hl-kurang">Kurang</span> sehingga menjadi fokus utama intervensi sekolah'
        )
        if kurang_list:
            daftar = "; ".join(f'<span class="hl-kurang">{i + 1}) {nama}</span>' + (f" (skor {skor})" if skor is not None else "") for i, (nama, skor) in enumerate(kurang_list))
            kesimpulan += f", yaitu: {daftar}."
    elif jpk["Sedang"] > 0:
        kesimpulan += " Indikator yang masih berpredikat Sedang perlu diperkuat agar dapat mencapai predikat Baik."
    else:
        kesimpulan += " Capaian mutu pendidikan secara keseluruhan sudah berjalan dengan baik."

    ringkasan = {"total_indikator": jumlah_utama, "jumlah_per_kategori": jpk, "rata_rata_skor_2025": rata2, "indikator_naik": naik, "indikator_turun": turun}

    return RaporResponse(
        judul=judul, capaian=grouped["Capaian"], baik=grouped["Baik"], sedang=grouped["Sedang"], kurang=grouped["Kurang"],
        ringkasan=ringkasan, kesimpulan=kesimpulan,
    )

def _generate_rekomendasi(result):
    rekomendasi = []
    all_kurang = []
    for item in result.kurang:
        all_kurang.append(item)
        for sub in item.sub_indikator: 
            all_kurang.append(sub)

    idx = 1
    for item in all_kurang:
        kode_no = str(item.no).strip().upper()
        nama_capaian = str(item.capaian).lower()
        
        # Cari matching berdasarkan kode indikator atau nama capaian
        matched_data = None
        for key_kode, val_map in KODE_BENAHI_MAP.items():
            if kode_no.startswith(key_kode) or key_kode.lower() in nama_capaian:
                matched_data = val_map
                break
        
        skor_fmt = f"{item.skor_2025:.2f}" if item.skor_2025 is not None else "-"
        
        if matched_data:
            for b in matched_data["benahi"]:
                rekomendasi.append({
                    "no": idx,
                    "kode_indikator": item.no,
                    "indikator": item.capaian,
                    "skor": skor_fmt,
                    "identifikasi_masalah": f"Capaian {item.capaian} belum optimal",
                    "refleksi_akar_masalah": matched_data["akar"],
                    "kode_pbd": b["kode_pbd"],
                    "kegiatan_pembenahan": b["kegiatan_pbd"],
                    "kode_arkas": b["kode_arkas"],
                    "kegiatan_arkas": b["kegiatan_arkas"],
                    "estimasi_anggaran": "Disesuaikan ARKAS"
                })
                idx += 1
        else:
            rekomendasi.append({
                "no": idx,
                "kode_indikator": item.no,
                "indikator": item.capaian,
                "skor": skor_fmt,
                "identifikasi_masalah": f"Capaian {item.capaian} masih rendah",
                "refleksi_akar_masalah": "Perlu analisis berbasis data bersama tim PBD",
                "kode_pbd": "PBD-GEN-01",
                "kegiatan_pembenahan": "Evaluasi dan penyusunan program pembenahan",
                "kode_arkas": "03.01.01",
                "kegiatan_arkas": "Penyusunan RKT / RKAS",
                "estimasi_anggaran": "Disesuaikan"
            })
            idx += 1

    return rekomendasi

# ============ ENDPOINTS ============
@app.post("/api/analyze", response_model=RaporResponse)
async def analyze(file: UploadFile = File(...), nama_sekolah: Optional[str] = Form(None)):
    nama = file.filename or ""
    ext = "." + nama.rsplit(".", 1)[-1].lower() if "." in nama else ""
    if ext not in ALLOWED_EXT:
        raise HTTPException(400, f"Format tidak didukung: {ext}")
    
    content = await file.read()
    if not content:
        raise HTTPException(400, "File kosong.")

    try:
        if ext in (".xlsx", ".xls"):
            df, judul = parse_excel(content)
        elif ext == ".pdf":
            df, judul = parse_pdf(content), ""
        else:
            df, judul = parse_docx(content), ""
    except Exception as e:
        raise HTTPException(422, f"Gagal memparse file: {e}")

    if df.empty:
        raise HTTPException(422, "Tidak ada data yang bisa diekstrak dari file.")

    meta = _ekstrak_meta(judul, nama)
    if nama_sekolah:
        meta['nama_sekolah'] = nama_sekolah
    judul_akhir = f"Analisis Rapor Pendidikan {meta['nama_sekolah']} Tahun {meta['tahun']}"

    try:
        hasil = _analisis(df, judul_akhir)
        
        # Simpan ke Database SQLite
        r_data = hasil.ringkasan
        jpk = r_data.get("jumlah_per_kategori", {})

        query = (
            "INSERT INTO history ("
            "nama_sekolah, tahun, filename, total_indikator, "
            "rata_rata_skor, indikator_naik, indikator_turun, "
            "jumlah_baik, jumlah_sedang, jumlah_kurang, data"
            ") VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)"
        )

        params = (
            meta['nama_sekolah'],
            meta['tahun'],
            nama,
            r_data.get("total_indikator", 0),
            r_data.get("rata_rata_skor_2025") or 0.0,
            r_data.get("indikator_naik", 0),
            r_data.get("indikator_turun", 0),
            jpk.get("Baik", 0),
            jpk.get("Sedang", 0),
            jpk.get("Kurang", 0),
           json.dumps(hasil.model_dump())
        )

        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        c.execute(query, params)
        conn.commit()
        conn.close()

        return hasil
    except Exception as e:
        raise HTTPException(422, f"Analisis gagal: {e}")

@app.post("/api/recommendations")
async def recommendations(data: dict):
    try:
        resp = RaporResponse(**data)
        return {"rekomendasi": _generate_rekomendasi(resp)}
    except Exception as e:
        raise HTTPException(500, f"Gagal generate rekomendasi: {e}")

@app.post("/api/compare")
async def compare(files: List[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(400, "Upload minimal 2 file.")
    results = []
    for f in files:
        try:
            content = await f.read()
            ext = "." + f.filename.rsplit(".", 1)[-1].lower() if "." in f.filename else ""
            if ext in (".xlsx", ".xls"): df, judul = parse_excel(content)
            elif ext == ".pdf": df, judul = parse_pdf(content), ""
            elif ext == ".docx": df, judul = parse_docx(content), ""
            else: continue
            
            if "skor_2025" not in df.columns: continue
            meta = _ekstrak_meta(judul, f.filename)
            skors = df["skor_2025"].apply(_clean_angka).dropna()
            results.append({"sekolah": meta["nama_sekolah"], "avg_score": round(float(skors.mean()), 2) if len(skors) > 0 else 0})
        except Exception:
            continue
    return results

@app.post("/api/export-pdf")
async def export_pdf(data: dict):
    try:
        from fpdf import FPDF
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=16)
        pdf.cell(0, 12, data.get("judul", "Analisis Rapor Pendidikan"), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(8)
        kesimpulan = re.sub(r"<[^>]+>", "", data.get("kesimpulan", ""))
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 8, kesimpulan)
        pdf.ln(8)
        for kelompok in ["baik", "sedang", "kurang"]:
            items = data.get(kelompok, [])
            if items:
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, f"KELOMPOK {kelompok.upper()} ({len(items)})", new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Arial", size=10)
                for item in items:
                    pdf.cell(0, 7, f"- {item.get('no','')} {item.get('capaian','')} (Skor: {item.get('skor_2025','-')})", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(4)
        return StreamingResponse(io.BytesIO(pdf.output()), media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=rapor_pendidikan.pdf"})
    except Exception as e:
        raise HTTPException(500, f"Gagal generate PDF: {e}")

# ============ ENDPOINT HISTORY (SQLite CRUD) ============
@app.get("/api/history")
async def get_history():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("""
        SELECT id, nama_sekolah, tahun, filename, total_indikator, 
               rata_rata_skor, created_at 
        FROM history 
        ORDER BY created_at DESC
    """)
    rows = [dict(r) for r in c.fetchall()]
    conn.close()
    return rows

@app.get("/api/history/{hid}")
async def get_history_detail(hid: int):
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    c = conn.cursor()
    c.execute("SELECT * FROM history WHERE id = ?", (hid,))
    row = c.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Riwayat tidak ditemukan.")
    
    result = dict(row)
    try:
        result["data"] = json.loads(result["data"])
    except Exception:
        result["data"] = {}
        
    return result

@app.delete("/api/history/{hid}")
async def delete_history(hid: int):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("DELETE FROM history WHERE id = ?", (hid,))
    deleted = c.rowcount
    conn.commit()
    conn.close()
    
    if deleted == 0:
        raise HTTPException(status_code=404, detail="Riwayat tidak ditemukan.")
        
    return {"status": "ok", "message": "Riwayat berhasil dihapus."}

@app.get("/api/health")
def health():
    return {"status": "ok"}
@app.post("/api/export-recommendations-excel")
async def export_recommendations_excel(data: dict):
    """Mengunduh data Rekomendasi PBD & ARKAS dalam format Excel (.xlsx)."""
    try:
        resp = RaporResponse(**data)
        rekom_list = _generate_rekomendasi(resp)
        
        # Konversi data rekomendasi ke DataFrame Pandas
        df = pd.DataFrame(rekom_list)
        
        # Format nama kolom agar rapi saat dibuka di Excel
        df = df.rename(columns={
            "no": "No",
            "kode_indikator": "Kode Indikator",
            "indikator": "Indikator / Capaian",
            "skor": "Skor 2025",
            "identifikasi_masalah": "Identifikasi Masalah",
            "refleksi_akar_masalah": "Refleksi Akar Masalah",
            "kode_pbd": "Kode PBD",
            "kegiatan_pembenahan": "Program Pembenahan (PBD)",
            "kode_arkas": "Kode ARKAS",
            "kegiatan_arkas": "Kegiatan ARKAS",
            "estimasi_anggaran": "Estimasi Anggaran"
        })

        output = io.BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="Rekomendasi PBD-ARKAS")
        output.seek(0)

        filename = "Rekomendasi_PBD_ARKAS.xlsx"
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(500, f"Gagal export Rekomendasi Excel: {e}")


@app.post("/api/export-recommendations-pdf")
async def export_recommendations_pdf(data: dict):
    """Mengunduh data Rekomendasi PBD & ARKAS dalam format PDF."""
    try:
        from fpdf import FPDF
        
        resp = RaporResponse(**data)
        rekom_list = _generate_rekomendasi(resp)
        
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Judul Dokumen
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "REKOMENDASI PROGRAM BENAHI (PBD) & ARKAS", new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(5)

        if not rekom_list:
            pdf.set_font("Arial", size=11)
            pdf.cell(0, 10, "Tidak ada rekomendasi program yang diperlukan.", new_x="LMARGIN", new_y="NEXT")
        else:
            for item in rekom_list:
                pdf.set_font("Arial", "B", 11)
                pdf.set_text_color(180, 0, 0) # Warna merah tua untuk header item
                pdf.cell(0, 7, f"[{item['kode_indikator']}] {item['indikator']} (Skor: {item['skor']})", new_x="LMARGIN", new_y="NEXT")
                
                pdf.set_font("Arial", size=10)
                pdf.set_text_color(0, 0, 0)
                
                # Menulis detail rekomendasi
                txt_akar = f"  - Akar Masalah: {item['refleksi_akar_masalah']}"
                txt_pbd = f"  - Program PBD ({item['kode_pbd']}): {item['kegiatan_pembenahan']}"
                txt_arkas = f"  - Kegiatan ARKAS ({item['kode_arkas']}): {item['kegiatan_arkas']}"
                
                pdf.multi_cell(0, 6, txt_akar)
                pdf.multi_cell(0, 6, txt_pbd)
                pdf.multi_cell(0, 6, txt_arkas)
                pdf.ln(4)

        output = io.BytesIO(pdf.output())
        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=Rekomendasi_PBD_ARKAS.pdf"}
        )
    except Exception as e:
        raise HTTPException(500, f"Gagal export Rekomendasi PDF: {e}")
# ============ SERVE HTML ============
_BACKEND_DIR = os.path.dirname(os.path.abspath(__file__))
_HTML_DIR = os.path.join(_BACKEND_DIR, "html")
_index_file = os.path.join(_HTML_DIR, "index.html")

if os.path.isfile(_index_file):
    @app.get("/", include_in_schema=False)
    async def root():
        return FileResponse(_index_file)
else:
    print(f"WARNING: index.html tidak ditemukan di {_HTML_DIR}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
    
    
    # ============ ENDPOINT EXPORT REKOMENDASI ============

